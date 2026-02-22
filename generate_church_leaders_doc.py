"""
Generate a professional Word document describing the Church Attendance System for church leaders.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_with_color(doc, text, level, color_hex=None):
    """Add a heading with optional color."""
    heading = doc.add_heading(text, level=level)
    if color_hex and level == 1:
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*bytes.fromhex(color_hex))
    return heading

def shade_paragraph(paragraph, color_hex):
    """Add background color to a paragraph."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    paragraph._element.get_or_add_pPr().append(shading_elm)

def create_church_leaders_document():
    """Create the Word document."""
    doc = Document()
    
    # Set up document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title Page
    title = doc.add_heading('Church Attendance Tracking System', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title.runs:
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(31, 78, 121)  # Dark blue
    
    subtitle = doc.add_paragraph('A Modern Solution for Church Leadership')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(16)
        run.font.italic = True
        run.font.color.rgb = RGBColor(70, 130, 180)
    
    doc.add_paragraph()  # Spacing
    
    date_para = doc.add_paragraph('February 2026')
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_page_break()
    
    # Executive Summary
    add_heading_with_color(doc, 'Executive Summary', 1, '1F4E79')
    
    summary_text = """The Church Attendance Tracking System is a modern, user-friendly solution designed to streamline how your church manages member attendance. This system enables church leaders to efficiently track attendance, understand congregation engagement patterns, and make data-driven decisions about ministry planning and growth."""
    
    doc.add_paragraph(summary_text)
    
    doc.add_paragraph()
    
    # Key Benefits
    add_heading_with_color(doc, 'Key Benefits for Church Leadership', 1, '1F4E79')
    
    benefits = [
        ('Increased Efficiency', 'Reduce attendance tracking time from manual entry to seconds with QR code scanning or quick manual input. Free up staff to focus on ministry.'),
        ('Better Member Insights', 'Access detailed attendance analytics and engagement patterns. Identify inactive members and opportunities for pastoral outreach.'),
        ('Data-Driven Decisions', 'Make informed decisions about service times, worship styles, and ministry programs based on real attendance data and trends.'),
        ('Improved Accessibility', 'Easy-to-use mobile and desktop interfaces that require minimal training. Works in any browser without installing applications.'),
        ('Enhanced Member Care', 'Quickly identify members who may need support or mentoring. Track attendance for accountability groups and small group ministries.'),
        ('Cost Effective', 'Eliminate paper forms and manual record-keeping. Reduce administrative overhead and improve accuracy.'),
    ]
    
    for title, description in benefits:
        p = doc.add_paragraph(style='List Bullet')
        p_format = p.paragraph_format
        p_format.left_indent = Inches(0.25)
        
        run = p.add_run(title + ': ')
        run.bold = True
        run.font.color.rgb = RGBColor(31, 78, 121)
        p.add_run(description)
    
    doc.add_page_break()
    
    # Core Features
    add_heading_with_color(doc, 'Core Features', 1, '1F4E79')
    
    doc.add_heading('Member Management', level=2)
    doc.add_paragraph('Maintain complete member profiles with contact information, demographics, and group assignments. Easily add new members and manage member status.', style='Normal')
    
    doc.add_heading('Service Organization', level=2)
    doc.add_paragraph('Organize and manage church services and events. Schedule different service times and special occasions, and track attendance for each event.', style='Normal')
    
    doc.add_heading('QR Code Check-In', level=2)
    doc.add_paragraph('Modern check-in experience using QR codes for each member. Simply scan member QR codes for instant attendance recording, or enter member names manually.', style='Normal')
    
    doc.add_heading('Real-Time Reports', level=2)
    doc.add_paragraph('Access comprehensive attendance analytics at a glance. View attendance trends, identify patterns, and generate detailed reports for ministry planning.', style='Normal')
    
    doc.add_heading('Admin Dashboard', level=2)
    doc.add_paragraph('Full administrative control with secure authentication. Manage system settings, user roles, and generate custom reports for leadership meetings.', style='Normal')
    
    doc.add_page_break()
    
    # Use Cases
    add_heading_with_color(doc, 'Real-World Applications', 1, '1F4E79')
    
    use_cases = [
        {
            'title': 'Service Attendance Tracking',
            'description': 'Track attendance for Sunday services, Wednesday night services, special events, and seasonal services. Understand sanctuary attendance trends throughout the year.'
        },
        {
            'title': 'Small Group Management',
            'description': 'Monitor attendance in small groups, Bible studies, youth groups, and ministry teams. Identify members who may need follow-up or prayer support.'
        },
        {
            'title': 'Newcomer Integration',
            'description': 'Track first-time visitors and repeat visitors. Identify which newcomers are becoming integrated into the church community.'
        },
        {
            'title': 'Ministry Effectiveness',
            'description': 'Evaluate the effectiveness of worship services, sermon topics, and special programs based on attendance data.'
        },
        {
            'title': 'Leadership Planning',
            'description': 'Make informed decisions about facility needs, staffing levels, and ministry expansion based on actual attendance trends.'
        },
    ]
    
    for use_case in use_cases:
        doc.add_heading(use_case['title'], level=2)
        doc.add_paragraph(use_case['description'])
    
    doc.add_page_break()
    
    # How It Works
    add_heading_with_color(doc, 'How It Works', 1, '1F4E79')
    
    doc.add_heading('For Members', level=2)
    steps_members = [
        'Members receive a unique QR code (physical card or digital)',
        'At check-in, simply scan the QR code or provide your name',
        'Attendance is recorded instantly in the system',
        'No lengthy sign-up sheets or forms needed'
    ]
    for i, step in enumerate(steps_members, 1):
        doc.add_paragraph(step, style='List Number')
    
    doc.add_heading('For Leadership', level=2)
    steps_leaders = [
        'Access the admin dashboard to view real-time attendance data',
        'Generate detailed reports for specific services, dates, or members',
        'Identify trends and patterns in attendance',
        'Export data for ministry planning and budgeting decisions',
        'Manage members, services, and system settings securely'
    ]
    for i, step in enumerate(steps_leaders, 1):
        doc.add_paragraph(step, style='List Number')
    
    doc.add_page_break()
    
    # Technical Assurance
    add_heading_with_color(doc, 'Technical Reliability & Security', 1, '1F4E79')
    
    doc.add_paragraph('This system is built with enterprise-grade technology:')
    
    features = [
        'Built on Python Django and React - proven, stable technologies used by major organizations',
        'Responsive design works on computers, tablets, and smartphones',
        'Works online and can handle large congregations',
        'Secure authentication protects member data and leadership information',
        'Regular backups ensure no data is ever lost',
        'Easy to deploy and maintain with minimal IT infrastructure',
    ]
    
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_page_break()
    
    # Implementation
    add_heading_with_color(doc, 'Getting Started', 1, '1F4E79')
    
    doc.add_paragraph(
        'Implementation is straightforward and can be completed in phases:'
    )
    
    phases = [
        {
            'phase': 'Phase 1: Setup',
            'description': 'Install the system, configure basic settings, and import member list'
        },
        {
            'phase': 'Phase 2: Training',
            'description': 'Train volunteer staff and leadership on how to use the system'
        },
        {
            'phase': 'Phase 3: Launch',
            'description': 'Begin using for one service, then expand to all services and events'
        },
        {
            'phase': 'Phase 4: Optimization',
            'description': 'Fine-tune processes and begin leveraging data for ministry decisions'
        },
    ]
    
    for phase_info in phases:
        doc.add_heading(phase_info['phase'], level=2)
        doc.add_paragraph(phase_info['description'])
    
    doc.add_page_break()
    
    # ROI and Impact
    add_heading_with_color(doc, 'Return on Investment', 1, '1F4E79')
    
    doc.add_paragraph('Measurable benefits of implementing this system:')
    
    roi_items = [
        'Save staff time previously spent on manual attendance tracking (5-10 hours per month)',
        'Improve accuracy of attendance records by eliminating human error',
        'Identify growth opportunities through data-driven insights',
        'Enhance member engagement through personalized pastoral care',
        'Support membership discipline and accountability in a loving way',
        'Create historical records for annual reports and ministry reviews',
        'Reduce paper usage and administrative costs',
    ]
    
    for item in roi_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # Conclusion
    add_heading_with_color(doc, 'Conclusion', 1, '1F4E79')
    
    conclusion = """The Church Attendance Tracking System is a powerful tool that helps church leadership fulfill their calling more effectively. By automating attendance tracking and providing actionable insights, this system frees leaders to focus on what matters most: pastoral care, spiritual growth, and building a thriving church community.

Whether you're a small church of fifty members or a large congregation of thousands, this system scales to meet your needs. It's designed with church leaders in mind, combining simplicity for users with powerful analytics for decision-makers.

We invite you to explore this system and see how it can enhance your church's ministry and growth."""
    
    doc.add_paragraph(conclusion)
    
    doc.add_page_break()
    
    # Contact/Support
    add_heading_with_color(doc, 'Next Steps', 1, '1F4E79')
    
    next_steps = [
        'Schedule a demonstration to see the system in action',
        'Review the implementation timeline with your team',
        'Discuss data privacy and security requirements',
        'Plan for staff training and member communication',
        'Begin the setup and deployment process',
    ]
    
    for step in next_steps:
        doc.add_paragraph(step, style='List Number')
    
    # Save document
    output_path = r'c:\Users\DELL\Desktop\Church_Attendance\Church_Attendance_System_for_Leaders.docx'
    doc.save(output_path)
    print(f'✓ Document created successfully: {output_path}')
    return output_path

if __name__ == '__main__':
    create_church_leaders_document()
